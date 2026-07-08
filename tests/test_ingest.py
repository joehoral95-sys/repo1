from deckstudio.ingest.digest import ingest_file


def test_docx_digest(tmp_path):
    import docx

    doc = docx.Document()
    doc.add_heading("Quarterly Results", level=1)
    doc.add_paragraph("Membership grew 2.1 percent.")
    doc.add_heading("Risks", level=2)
    doc.add_paragraph("Prelim exams dipped.", style="List Bullet")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Members"
    table.rows[1].cells[1].text = "33412"
    src = tmp_path / "report.docx"
    doc.save(str(src))

    out = ingest_file(src, tmp_path / "sources")
    text = out.read_text(encoding="utf-8")
    assert "## Quarterly Results" in text
    assert "### Risks" in text
    assert "- Prelim exams dipped." in text
    assert "| Members" in text and "33412" in text


def test_xlsx_digest(tmp_path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Exams"
    ws.append(["Exam", "Q1", "Q2"])
    ws.append(["P", 2410, 2180])
    ws.append(["FM", 2205, 1990])
    src = tmp_path / "exams.xlsx"
    wb.save(str(src))

    out = ingest_file(src, tmp_path / "sources")
    text = out.read_text(encoding="utf-8")
    assert "## Sheet: Exams" in text
    assert "2410" in text and "1990" in text


def test_txt_digest(tmp_path):
    src = tmp_path / "notes.txt"
    src.write_text("audience is all staff\n", encoding="utf-8")
    out = ingest_file(src, tmp_path / "sources")
    assert "all staff" in out.read_text(encoding="utf-8")


def test_unknown_extension_message(tmp_path):
    src = tmp_path / "weird.xyz"
    src.write_text("x", encoding="utf-8")
    try:
        ingest_file(src, tmp_path / "sources")
        raise AssertionError("should have raised")
    except ValueError as e:
        assert "deckstudio extract" in str(e)
